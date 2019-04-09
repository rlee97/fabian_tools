from docx import Document
from pprint import pprint

# NOTE UPDATE THESE VARIABLES TO MAKE SURE THAT IT WORKS PROPERLY NOTE UPDATE THESE VARIABLES TO MAKE SURE THAT IT WORKS
# Enter the location for the git stats file
stats_file_location = "C:\\Users\\user.name\\Desktop\\Directory\\fabian-gui\\FabianEvo\\stats.txt"
# Enter the location of the document you want to update
document_file_location = "C:\\Users\\user.name\\Desktop\\Directory\\Code Review Report - GUI.docx"
# Enter the location you want the document to be updated. This can be the same as the document the same as above
document_file_save_location = "C:\\Users\\user.name\\Desktop\\Directory\\Code Review Report - GUI.docx"
# Enter the number you want to start the table to start at
start_counter = 1
# Enter the table you want to start at
table_counter = 9
# NOTE UPDATE THESE VARIABLES TO MAKE SURE THAT IT WORKS PROPERLY NOTE UPDATE THESE VARIABLES TO MAKE SURE THAT IT WORKS


class Table_Automation_Write:

    # Initialization function for the automation reader command
    def __init__(self, input_stats_location, input_document_location, input_document_save_location, input_start_counter, input_table_counter):
        self.stats_location = input_stats_location
        self.document_location = input_document_location
        self.document_save_location = input_document_save_location
        self.start_counter = input_start_counter
        self.table_counter = input_table_counter

    # This function will automate the entire fill in process from the stats file to the document
    # NOTE: This function only appends to the table not inserts
    def automate(self):
        # Document opening for output
        document = Document(self.document_location)
        tables = document.tables

        # Reading in the stats file
        stats_file = open(self.stats_location, "r")
        counter = self.start_counter
        index = 0
        buffer = ""

        # This will go through the stats file and find the files with the associated ending tags in them
        for line in stats_file:
            if((".c " in line) or (".C " in line) or (".h " in line) or (".H " in line) or (".cpp " in line) or (".rc " in line) or (".rc2 " in line)):
                # Table at 4 correspond to the table we want to start writing into
                index = line.find("|")
                buffer = line[:index]
                row_cells = tables[self.table_counter].add_row().cells
                # The row 0 cells account for the numbers portion of the table
                row_cells[0].text = str(counter)
                # The row 1 cells account for the text portion of the table
                row_cells[1].text = buffer
                # This increments the counter for the table to insert in the correct one
                counter += 1
            else:
                # Nothing needs to be done here since the file is not cared about
                pass

        # Saves the document
        document.save(self.document_save_location)

    # This function will only fill in the numbers on the table
    def fill_in_numbers(self, input_counter, input_doc):
        document = Document(input_doc)
        # Saves the document tables into the variable tables
        tables = document.tables
        print("Tables:")
        pprint(tables)

        counter = input_counter

        table = tables[table_counter]
        flag = False
        for row in table.rows:
            # This passes the first row in the table [No, Filename]
            if(flag == False):
                assign_row = row.cells
                assign_row[0].text = ""
                flag = True
            else:
                # This puts the counter into the appropriate cell
                assign_row = row.cells
                assign_row[0].text = str(counter)
                counter += 1

        # Saves the document
        document.save(document_file_save_location)


# This function will
def main():
    table_write = Table_Automation_Write(stats_file_location, document_file_location, 
                                         document_file_save_location, start_counter, table_counter)
    # Automates the fill in process
    table_write.automate()
    # This commented in will fill in the table with the correct counters
    # table_write.fill_in_numbers(1, document_file_location)


if __name__ == '__main__':
    main()
